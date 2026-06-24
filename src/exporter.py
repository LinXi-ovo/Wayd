# Export functionality for WAYD records - DOCX and PDF
import os
from datetime import datetime
import tkinter as tk
from tkinter import filedialog, messagebox
import sqlite3
from PIL import Image
import io

DB_FILE = "whatido.db"

def get_all_filtered_records(date_filter=None, keyword=None):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    query = "SELECT id, timestamp, doing, next_plan, screenshot_path FROM records WHERE 1=1"
    params = []
    if date_filter:
        query += " AND date(timestamp) = ?"
        params.append(date_filter)
    if keyword:
        query += " AND (doing LIKE ? OR next_plan LIKE ?)"
        like = f"%{keyword}%"
        params.extend([like, like])
    query += " ORDER BY timestamp DESC"
    c.execute(query, params)
    rows = c.fetchall()
    conn.close()
    return rows


def _resize_image_for_docx(image_path, max_width=1200, quality=75):
    """Compress image for DOCX embedding, return BytesIO"""
    try:
        img = Image.open(image_path)
        if img.width > max_width:
            ratio = max_width / img.width
            new_size = (max_width, int(img.height * ratio))
            img = img.resize(new_size, Image.LANCZOS)
        if img.mode == "RGBA":
            bg = Image.new("RGB", img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[3])
            img = bg
        elif img.mode != "RGB":
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=quality)
        buf.seek(0)
        return buf
    except Exception as e:
        print(f"[exporter] compress fail: {e}")
        return None

def export_docx(date_filter, keyword, compressed=False):
    """Export records as DOCX with embedded screenshots"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        messagebox.showerror(
            "缺少依赖",
            "请先安装 python-docx：\nuv add python-docx"
        )
        return

    records = get_all_filtered_records(date_filter, keyword)
    if not records:
        messagebox.showwarning("提示", "当前筛选条件下没有记录可导出")
        return

    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word 文档", "*.docx")],
        initialfile=f"WAYD_报告_{datetime.now():%Y%m%d_%H%M%S}.docx",
        title="导出为 DOCX"
    )
    if not file_path:
        return

    try:
        doc = Document()
        title = doc.add_heading("WAYD - 时间追踪报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"导出时间：{datetime.now():%Y-%m-%d %H:%M:%S}")
        doc.add_paragraph(f"记录总数：{len(records)} 条")
        if date_filter:
            doc.add_paragraph(f"日期筛选：{date_filter}")
        if keyword:
            doc.add_paragraph(f"关键词筛选：{keyword}")

        doc.add_paragraph("─" * 60)

        for i, (rid, ts, doing, next_plan, screenshot) in enumerate(records, 1):
            doc.add_heading(f"记录 #{i}  (ID: {rid})", level=2)
            doc.add_paragraph(f"时间：{ts}")
            p = doc.add_paragraph()
            p.add_run("在干嘛：").bold = True
            p.add_run(doing)
            p = doc.add_paragraph()
            p.add_run("下一步：").bold = True
            p.add_run(next_plan)
            if screenshot and os.path.exists(screenshot):
                try:
                    if compressed:
                        img_data = _resize_image_for_docx(screenshot)
                        if img_data:
                            doc.add_picture(img_data, width=Inches(5.5))
                        else:
                            doc.add_picture(screenshot, width=Inches(5.5))
                    else:
                        doc.add_picture(screenshot, width=Inches(5.5))
                except Exception:
                    doc.add_paragraph("[截图嵌入失败]")
            else:
                doc.add_paragraph("截图：无")
            doc.add_paragraph("")

        doc.save(file_path)
        messagebox.showinfo("成功", f"已导出至：\n{file_path}")
    except Exception as e:
        messagebox.showerror("导出失败", f"DOCX 导出出错：\n{e}")

def _get_pdf_font():
    """Find available Chinese font on system for PDF export"""
    candidates = [
        (r"C:\Windows\Fonts\msyh.ttc", 0),
        (r"C:\Windows\Fonts\msyhbd.ttc", 0),
        (r"C:\Windows\Fonts\simsun.ttc", 0),
        (r"C:\Windows\Fonts\simhei.ttf", None),
    ]
    for path, idx in candidates:
        if os.path.exists(path):
            return path, idx
    return None, None

def export_pdf(date_filter, keyword):
    """Export records as PDF with embedded screenshots"""
    try:
        from fpdf import FPDF
    except ImportError:
        messagebox.showerror(
            "缺少依赖",
            "请先安装 fpdf2：\nuv add fpdf2"
        )
        return

    records = get_all_filtered_records(date_filter, keyword)
    if not records:
        messagebox.showwarning("提示", "当前筛选条件下没有记录可导出")
        return

    file_path = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("PDF 文档", "*.pdf")],
        initialfile=f"WAYD_报告_{datetime.now():%Y%m%d_%H%M%S}.pdf",
        title="导出为 PDF"
    )
    if not file_path:
        return

    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)

        font_path, font_idx = _get_pdf_font()
        has_cjk = False
        if font_path:
            try:
                kwargs = {"ttc_font_index": font_idx} if font_idx is not None else {}
                pdf.add_font("CJK", "", font_path, uni=True, **kwargs)
                pdf.add_font("CJK", "B", font_path, uni=True, **kwargs)
                has_cjk = True
            except Exception:
                pass

        # Title page
        pdf.add_page()
        if has_cjk:
            pdf.set_font("CJK", "B", 20)
        else:
            pdf.set_font("Helvetica", "B", 20)
        pdf.cell(0, 15, "WAYD - Time Tracking Report", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(10)

        if has_cjk:
            pdf.set_font("CJK", "", 10)
        else:
            pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 7, f"Export: {datetime.now():%Y-%m-%d %H:%M:%S}", new_x="LMARGIN", new_y="NEXT")
        pdf.cell(0, 7, f"Records: {len(records)}", new_x="LMARGIN", new_y="NEXT")
        if date_filter:
            pdf.cell(0, 7, f"Date: {date_filter}", new_x="LMARGIN", new_y="NEXT")
        if keyword:
            pdf.cell(0, 7, f"Keyword: {keyword}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(5)

        for i, (rid, ts, doing, next_plan, screenshot) in enumerate(records, 1):
            pdf.add_page()
            if has_cjk:
                pdf.set_font("CJK", "B", 14)
            else:
                pdf.set_font("Helvetica", "B", 14)
            pdf.cell(0, 10, f"Record #{i} (ID: {rid})", new_x="LMARGIN", new_y="NEXT")

            if has_cjk:
                pdf.set_font("CJK", "", 10)
            else:
                pdf.set_font("Helvetica", "", 10)
            pdf.cell(0, 7, f"Time: {ts}", new_x="LMARGIN", new_y="NEXT")
            pdf.multi_cell(0, 7, f"Doing: {doing}")
            pdf.multi_cell(0, 7, f"Next: {next_plan}")

            if screenshot and os.path.exists(screenshot):
                try:
                    pdf.image(screenshot, w=150)
                except Exception:
                    pdf.cell(0, 7, "[screenshot embed failed]", new_x="LMARGIN", new_y="NEXT")
            else:
                pdf.cell(0, 7, "[no screenshot]", new_x="LMARGIN", new_y="NEXT")

        pdf.output(file_path)
        messagebox.showinfo("成功", f"已导出至：\n{file_path}")
    except Exception as e:
        messagebox.showerror("导出失败", f"PDF 导出出错：\n{e}")
